import fitz
from docx import Document
import pandas as pd
import os


#define parsing functions


#PDF
def parse_pdf(filepath):
    text=""
    docs = fitz.open(filepath)
    for page in docs:
        text+=page.get_text()
    return text.strip()

#WORD
def parse_word(filepath):
    docs=Document(filepath)
    text="\n".join([para.text for para in docs.paragraphs])
    # for para in docs.paragraphs:
    #     text+=para.text
    return text.strip()


#EXCEL
def parse_excel(filepath):
    df=pd.read_excel(filepath)
    return df.to_csv(index=False)



#process files in a directory
def process_files_in_directory(directory_path):
    parsed_data=[]

    for file in os.listdir(directory_path):
        file_path=os.path.join(directory_path,file)
        ext=os.path.splitext(file)[1].lower()
        
        print(f"\n📄 Processing: {file}")


        try:
            if ext == '.pdf':
                content=parse_pdf(file_path)
            
            elif ext == '.docx':
                content=parse_word(file_path)

            elif ext in ['.xls', '.xlsx']:
                content=parse_excel(file_path)

            else:
                print(f"❌ Unsupported file format: {ext}")
                continue

            parsed_data.append({
                "file_name": file,
                "extension": ext,
                "content": content
            })
            print(f"✅ Successfully processed: {file}")

        except Exception as e:
            print(f"❌ Error processing {file}: {e}")
    return parsed_data

#save extracted text to file:
def save_parsed_data(parsed_data, output_file):
    output_dir = "Outputs"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, output_file)
    with open(output_path, 'w', encoding='utf-8') as f:
        for data in parsed_data:
            f.write(f"=== {data['file_name']} ({data['extension']}) ===\n")
            f.write(data['content'])
            f.write("\n\n")

    print(f"\n📁 All parsed data saved successfully to: {output_path}")


#Main
if __name__ == "__main__":
    directory_path=input("Enter the directory path containing files to parse: ")
    output_file=input("Enter the output file name to save parsed data (e.g., parsed_output.txt): ")

    parsed_data=process_files_in_directory(directory_path)
    save_parsed_data(parsed_data, output_file)