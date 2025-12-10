# stage_1_parsing/word_parser.py

from docx import Document
import os
from typing import List, Tuple


def parse_word(file_path: str, session_id: str) -> Tuple[str, List[str]]:

    paragraphs = []
    saved_images = []

    images_dir = os.path.join("/tmp/outputs/images", "word_images", session_id)
    os.makedirs(images_dir, exist_ok=True)

    doc = Document(file_path)

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cleaned = ",".join([c.text.strip() for c in row.cells])
            if cleaned:
                paragraphs.append(cleaned)

    rels = doc.part.rels
    for rel in rels:
        rel_obj = rels[rel]

        if "image" in rel_obj.target_ref:
            img_bytes = rel_obj.target_part.blob
            img_name = os.path.basename(rel_obj.target_ref)
            img_path = os.path.join(images_dir, img_name)

            try:
                with open(img_path, "wb") as fh:
                    fh.write(img_bytes)
                saved_images.append(img_path)
            except Exception:
                continue

    full_text = "\n\n".join(paragraphs).strip()
    return full_text, saved_images
